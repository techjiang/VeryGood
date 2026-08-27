"""docx 抽取：段落（带标题层级）、表格、页眉页脚、超链接。"""
import re

from .common import merge_links, ooxml_external_links, text_urls

MAX_TABLE_ROWS = 500


def heading_level(p):
    """段落是否为标题及其层级（defect_check.py 的空章节检查复用同一判定）。"""
    try:
        name = (p.style.name or "").strip()
    except Exception:
        return None
    m = re.match(r"(?:Heading|标题)\s*(\d+)", name, re.I)
    return int(m.group(1)) if m else None


def extract(path):
    units, outline, notes = [], [], []
    try:
        import docx  # python-docx
    except Exception:
        return {
            "units": [],
            "links": merge_links(ooxml_external_links(path)),
            "outline": [],
            "meta": {},
            "notes": ["python-docx 不可用，docx 正文未抽取（pip install python-docx）"],
        }
    try:
        d = docx.Document(path)
    except Exception as e:
        # 打不开也把关系表里的外链交出去：链接可用性核验不应因正文失败而空手
        return {
            "units": [],
            "links": merge_links(ooxml_external_links(path)),
            "outline": [],
            "meta": {},
            "notes": [f"docx 打开失败: {e}"],
        }

    text_links = []
    for i, p in enumerate(d.paragraphs, 1):
        t = (p.text or "").strip()
        if not t:
            continue
        lvl = heading_level(p)
        anchor = f"docx:p{i}"
        if lvl:
            anchor += f"|H{lvl}"
            outline.append({"anchor": f"docx:p{i}", "label": t[:80], "level": lvl})
        units.append({"anchor": anchor, "text": t})
        text_links += text_urls(t, f"段落{i}")

    for ti, tbl in enumerate(d.tables, 1):
        rows = 0
        for ri, row in enumerate(tbl.rows, 1):
            if rows >= MAX_TABLE_ROWS:
                notes.append(f"表{ti} 超过 {MAX_TABLE_ROWS} 行，其余未抽取")
                break
            cells = [(c.text or "").replace("\n", " ").strip() for c in row.cells]
            if not any(cells):
                continue
            line = "\t".join(cells)
            units.append({"anchor": f"docx:t{ti}r{ri}", "text": line})
            text_links += text_urls(line, f"表{ti}第{ri}行")
            rows += 1

    for si, sec in enumerate(d.sections, 1):
        for kind, part in (("页眉", sec.header), ("页脚", sec.footer)):
            try:
                t = "\n".join((p.text or "").strip() for p in part.paragraphs).strip()
            except Exception:
                t = ""
            if t:
                units.append({"anchor": f"docx:s{si}{'h' if kind == '页眉' else 'f'}", "text": t})
                text_links += text_urls(t, f"第{si}节{kind}")

    return {
        "units": units,
        "links": merge_links(ooxml_external_links(path), text_links),
        "outline": outline,
        "meta": {"paragraphs": len(d.paragraphs), "tables": len(d.tables)},
        "notes": notes,
    }
